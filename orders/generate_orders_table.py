#!/usr/bin/env python3
"""
Process orders CSV file and generate a Word document with order details.
Each order is represented as a row with: Order no., Product, Size, Name, Number, Patch, Adds, Address
Uses translations.json for field translations (e.g., Hebrew to English).
"""

import csv
import sys
import os
import re
import json
import requests
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
import glob
import os
from datetime import datetime


# Global variable for product images base path
PRODUCTS_IMAGES_BASE = '../products/images'


def get_category_directory(category):
    """
    Map category name to its directory name.
    
    Args:
        category: Category extracted from product name
    
    Returns:
        Directory name for the category
    """
    category_mapping = {
        'אימונית': 'אימוניות',
        "ג'קט": "ג'קטים ומעילים",
        'מעיל רוח': "ג'קטים ומעילים",
        'חולצת': 'חולצות גברים',
        'חולצה ארוכה': 'חולצות גברים ארוכות',
        'חולצת נשים': 'חולצות נשים',
        'חליפת ילדים': 'חליפות ילדים',
        'מכנס': 'מכנסיים'
    }
    
    return category_mapping.get(category, category)


def extract_product_segments(product_name):
    """
    Extract category, team, type, and season from product name.
    
    Returns:
        tuple: (category, team, type, season) or (None, None, None, None) if parsing fails
    """
    if not product_name:
        return None, None, None, None
    
    # Define category patterns (order matters - check longer patterns first)
    categories = [
        'חולצת נשים',      # Women's shirt (check first - longer pattern)
        'חולצה ארוכה',      # Long sleeve shirt
        'חליפת ילדים',     # Kids kit
        'מעיל רוח',        # Windbreaker
        'חולצת',           # Shirt
        'מכנס',            # Shorts/Pants
        'אימונית',         # Training suit
        "ג'קט",           # Jacket
    ]
    
    # Find matching category
    category = None
    remaining = product_name
    
    for cat in categories:
        if product_name.startswith(cat):
            category = cat
            # Remove category from the beginning
            remaining = product_name[len(cat):].strip()
            break
    
    if not category:
        return None, None, None, None
    
    # Now extract team and type
    # Pattern: [team name] [type] [season]
    # Type patterns: בית, חוץ, השלישי, השלישית
    type_patterns = ['בית', 'חוץ', 'השלישי', 'השלישית']
    
    # Find the type keyword
    type_found = None
    type_index = -1
    
    for type_pattern in type_patterns:
        # Look for type pattern followed by space and year pattern
        pattern_with_space = f' {type_pattern} '
        idx = remaining.find(pattern_with_space)
        if idx != -1:
            type_found = type_pattern
            type_index = idx
            break
    
    if type_index == -1:
        # Try to find type at the end (before season)
        for type_pattern in type_patterns:
            if type_pattern in remaining:
                idx = remaining.find(f' {type_pattern}')
                if idx != -1:
                    type_found = type_pattern
                    type_index = idx
                    break
    
    if type_index == -1:
        return None, None, None, None
    
    # Team is everything before the type
    team = remaining[:type_index].strip()
    
    # Type is the keyword we found
    type_segment = type_found
    
    # Extract season from the end
    # Look for pattern like "2025/2026" or "2024/2025"
    season_match = re.search(r'(\d{4})/(\d{4})', remaining)
    season = None
    if season_match:
        year1 = season_match.group(1)
        year2 = season_match.group(2)
        season = f"{year1}-{year2}"
    
    return category, team, type_segment, season


def find_product_image(product_name, translations):
    """
    Find the first image (alphabetically) for a product based on its name.
    
    Args:
        product_name: Full product name
        translations: Translations dictionary (not used for images, kept for consistency)
    
    Returns:
        Path to image file or None if not found
    """
    category, team, type_segment, season = extract_product_segments(product_name)
    
    if not category or not team or not type_segment:
        return None
    
    # Get the directory name for the category
    category_dir = get_category_directory(category)
    
    # Construct path using Hebrew names directly
    # If season exists, prepend it to the path: ../products/images/2025-2026/[category]/[team]/[type]
    # Otherwise: ../products/images/[category]/[team]/[type]
    if season:
        image_dir = os.path.join(PRODUCTS_IMAGES_BASE, season, category_dir, team, type_segment)
    else:
        image_dir = os.path.join(PRODUCTS_IMAGES_BASE, category_dir, team, type_segment)
    
    # Normalize the path to handle escaped spaces and other issues
    image_dir = os.path.normpath(image_dir)
    
    # Check if directory exists
    if not os.path.exists(image_dir):
        print(f"Warning: Image directory not found: {image_dir}")
        return None
    
    # Find all images in directory (jpg, jpeg, png, gif, webp)
    image_extensions = ['*.jpg', '*.jpeg', '*.png', '*.gif', '*.webp', '*.JPG', '*.JPEG', '*.PNG']
    
    all_images = []
    for ext in image_extensions:
        # Use os.path.join which handles spaces properly
        pattern = os.path.join(image_dir, ext)
        images = glob.glob(pattern)
        all_images.extend(images)
    
    if all_images:
        # Sort alphabetically and return first
        all_images.sort()
        return all_images[0]
    
    return None


def insert_image_in_cell(cell, image_path, width_inches=2.0):
    """
    Insert an image into a table cell.
    Converts to JPG if the format is not recognized.
    
    Args:
        cell: The table cell to insert image into
        image_path: Path to the image file
        width_inches: Width of the image in inches
    """
    try:
        # Verify file exists and is readable
        if not os.path.exists(image_path):
            print(f"Warning: Image file does not exist: {image_path}")
            return False
        
        if not os.path.isfile(image_path):
            print(f"Warning: Path is not a file: {image_path}")
            return False
        
        # Check file size
        file_size = os.path.getsize(image_path)
        if file_size == 0:
            print(f"Warning: Image file is empty (0 bytes): {image_path}")
            return False
        
        # Try to insert the image directly first
        try:
            cell.text = ''
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(width_inches))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True
        except Exception as direct_error:
            # If direct insertion fails, try converting to JPG
            print(f"Info: Converting image to JPG format: {image_path}")
            
            try:
                from PIL import Image
                
                # Open and convert the image
                img = Image.open(image_path)
                
                # Convert to RGB if necessary (for PNG with transparency, etc.)
                if img.mode in ('RGBA', 'LA', 'P'):
                    # Create a white background
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'P':
                        img = img.convert('RGBA')
                    background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                    img = background
                elif img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # Save as temporary JPG
                temp_jpg_path = image_path + '.temp.jpg'
                img.save(temp_jpg_path, 'JPEG', quality=95)
                img.close()
                
                # Try inserting the converted image
                cell.text = ''
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(temp_jpg_path, width=Inches(width_inches))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Clean up temporary file
                try:
                    os.remove(temp_jpg_path)
                except:
                    pass
                
                print(f"Success: Converted and inserted image from {image_path}")
                return True
                
            except Exception as convert_error:
                print(f"Warning: Could not convert/insert image {image_path}: {type(convert_error).__name__}: {str(convert_error)}")
                return False
        
    except Exception as e:
        print(f"Warning: Could not insert image {image_path}: {type(e).__name__}: {str(e)}")
        return False


# Global variable for API key
ANTHROPIC_API_KEY = None
NO_TRANSLATE = False  # Flag to disable translation


def get_api_key():
    """Get Anthropic API key from cache, environment, or prompt user."""
    global ANTHROPIC_API_KEY
    
    # If NO_TRANSLATE flag is set, return None to skip translation
    if NO_TRANSLATE:
        return None
    
    if ANTHROPIC_API_KEY:
        return ANTHROPIC_API_KEY
    
    # Try to get from translations.json cache
    if os.path.exists(TRANSLATION_FILE):
        try:
            with open(TRANSLATION_FILE, 'r', encoding='utf-8') as f:
                translations = json.load(f)
                cached_key = translations.get('anthropic_api_key', '')
                if cached_key:
                    ANTHROPIC_API_KEY = cached_key
                    return cached_key
        except:
            pass
    
    # Try to get from environment variable
    api_key = os.environ.get('ANTHROPIC_API_KEY')
    
    if not api_key:
        # Prompt user
        print("\n" + "="*60)
        print("Anthropic API Key Required")
        print("="*60)
        print("To format addresses automatically, please provide your Anthropic API key.")
        print("You can get one from: https://console.anthropic.com/")
        print("Or set the ANTHROPIC_API_KEY environment variable.")
        api_key = input("Enter your API key (or press Enter to skip address formatting): ").strip()
        print("="*60 + "\n")
    
    if api_key:
        # Save to translations.json for future use
        try:
            translations = load_translations()
            translations['anthropic_api_key'] = api_key
            save_translations(translations)
            print("API key saved for future use.\n")
        except Exception as e:
            print(f"Warning: Could not save API key: {e}\n")
        
        ANTHROPIC_API_KEY = api_key
    
    return api_key


def validate_api_key(api_key):
    """
    Test if the API key is valid by making a simple request.
    Returns True if valid, False otherwise.
    """
    if not api_key:
        return False
    
    try:
        response = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "Content-Type": "application/json",
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01"
            },
            json={
                "model": "claude-sonnet-4-20250514",
                "max_tokens": 10,
                "messages": [
                    {"role": "user", "content": "Hi"}
                ]
            },
            timeout=10
        )
        
        return response.status_code == 200
    except:
        return False


def clear_invalid_api_key():
    """Remove invalid API key from cache and memory."""
    global ANTHROPIC_API_KEY
    ANTHROPIC_API_KEY = None
    
    try:
        translations = load_translations()
        if 'anthropic_api_key' in translations:
            del translations['anthropic_api_key']
            save_translations(translations)
            print("Invalid API key removed from cache.\n")
    except:
        pass


def set_cell_vertical_merge(cell, merge_type):
    """
    Set vertical merge for a cell.
    
    Args:
        cell: The cell to set merge on
        merge_type: 'restart' for first cell in merge, 'continue' for subsequent cells
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    vMerge = OxmlElement('w:vMerge')
    if merge_type == 'restart':
        vMerge.set(qn('w:val'), 'restart')
    # For 'continue', don't set the val attribute (empty vMerge element)
    tcPr.append(vMerge)


# Translation file path
TRANSLATION_FILE = 'translations.json'


def load_translations():
    """Load translations from JSON file."""
    if os.path.exists(TRANSLATION_FILE):
        with open(TRANSLATION_FILE, 'r', encoding='utf-8') as f:
            translations = json.load(f)
            # Ensure required sections exist
            if 'team_leagues' not in translations:
                translations['team_leagues'] = {}
            return translations
    # Return default structure with all sections
    return {
        "patch": {},
        "team_leagues": {}
    }


def save_translations(translations):
    """Save translations to JSON file."""
    with open(TRANSLATION_FILE, 'w', encoding='utf-8') as f:
        json.dump(translations, f, ensure_ascii=False, indent=2)


def translate_field(field_name, value, translations):
    """
    Translate a field value using the translations dictionary.
    If the field section exists but the value is not found, prompt user for translation.
    
    Args:
        field_name: The field section name (e.g., 'patch')
        value: The value to translate
        translations: The translations dictionary
    
    Returns:
        Translated value or original if no translation needed
    """
    if not value or not value.strip():
        return value
    
    value = value.strip()
    
    # Check if this field has a translation section
    if field_name not in translations:
        return value
    
    field_translations = translations[field_name]
    
    # Check if translation exists
    if value in field_translations:
        return field_translations[value]
    
    # Translation doesn't exist - prompt user
    print(f"\n{'='*60}")
    print(f"Translation needed for field: '{field_name}'")
    print(f"Hebrew text: {value}")
    user_translation = input(f"Enter English translation (or press Enter to keep as-is): ").strip()
    print(f"{'='*60}\n")
    
    # Save the translation
    if user_translation:
        translations[field_name][value] = user_translation
        save_translations(translations)
        return user_translation
    else:
        # User chose to keep as-is, save this choice
        translations[field_name][value] = value
        save_translations(translations)
        return value


def extract_team_name(product_name):
    """
    Extract team name from product name.
    Pattern: text between [חולצת|חולצה ארוכה|חולצת נשים|חליפת ילדים|מכנס] and [בית|חוץ|השלישי|השלישית]
    
    Args:
        product_name: The product name string
    
    Returns:
        Team name if found, None otherwise
    """
    if not product_name:
        return None
    
    # Define the prefix and suffix patterns
    prefix_pattern = r'(?:חולצת|חולצה ארוכה|חולצת נשים|חליפת ילדים|מכנס)\s+'
    suffix_pattern = r'\s+(?:בית|חוץ|השלישי|השלישית)'
    
    # Create full pattern to extract team name
    pattern = prefix_pattern + r'(.+?)' + suffix_pattern
    
    match = re.search(pattern, product_name)
    if match:
        return match.group(1).strip()
    
    return None


def get_league_patch_for_team(team_name, translations):
    """
    Get the league patch for a specific team.
    
    Args:
        team_name: The team name extracted from product
        translations: The translations dictionary
    
    Returns:
        League patch string (e.g., "LFP", "Premier League", etc.)
    """
    # Ensure team_leagues section exists
    if "team_leagues" not in translations:
        translations["team_leagues"] = {}
    
    team_leagues = translations["team_leagues"]
    
    # Check if team exists in mapping
    if team_name in team_leagues:
        return team_leagues[team_name]
    
    # Team not found - prompt user
    print(f"\n{'='*60}")
    print(f"League patch needed for team: '{team_name}'")
    print(f"Product contains team '{team_name}' but no league mapping exists.")
    user_patch = input(f"Enter league patch for this team (or press Enter for default 'League'): ").strip()
    print(f"{'='*60}\n")
    
    # Save the mapping
    if user_patch:
        translations["team_leagues"][team_name] = user_patch
        save_translations(translations)
        return user_patch
    else:
        # User chose default, save this choice
        translations["team_leagues"][team_name] = "League"
        save_translations(translations)
        return "League"


def find_latest_csv(directory='.'):
    """Find the most recent CSV file in the directory."""
    csv_files = glob.glob(os.path.join(directory, '*.csv'))
    if not csv_files:
        return None
    return max(csv_files, key=os.path.getctime)


def extract_option_value(options_text, key):
    """Extract value for a specific key from options string."""
    if not options_text:
        return ""
    
    # Pattern to match "key:value" where value can be in Hebrew or English
    pattern = rf'{re.escape(key)}:([^|]+)'
    match = re.search(pattern, options_text)
    if match:
        return match.group(1).strip()
    return ""


def extract_size(options_text):
    """Extract size from options string (מידה field)."""
    size = extract_option_value(options_text, "מידה")
    if size:
        # Extract just the size number/text before the colon or description
        size_match = re.match(r'([^:]+)', size)
        if size_match:
            return size_match.group(1).strip()
    return ""


def extract_patch(options_text, product_name, translations):
    """
    Extract patch information from options string (פאץ׳ field) and translate it.
    Special handling for "ליגה" - extract team name and lookup corresponding league patch.
    
    Args:
        options_text: The options/variant text containing patch info
        product_name: The product name for team extraction
        translations: The translations dictionary
    
    Returns:
        Translated patch string
    """
    patch = extract_option_value(options_text, "פאץ'") or extract_option_value(options_text, "פאץ׳")
    
    if not patch:
        return ""
    
    # Special handling for "ליגה"
    if patch == "ליגה":
        # Try to extract team name from product
        team_name = extract_team_name(product_name)
        
        if team_name:
            # Get league patch for this team
            return get_league_patch_for_team(team_name, translations)
        else:
            # Pattern doesn't match, return default "League"
            return "League"
    
    # For other patches, use normal translation
    return translate_field('patch', patch, translations)


def extract_adds(options_text):
    """
    Extract additional items from options string (הוסף field).
    Handles various add-on patterns and returns formatted English text.
    
    Patterns:
    - השלם לסט:הוסף מכנס -> + Shorts
    - השלם לסט:הוסף מכנס+גרביים -> + Shorts\n+ Socks
    - הוסף גרביים:כן -> + Socks
    - הוסף מכנסיים:כן -> + Pants
    """
    if not options_text:
        return ""
    
    # Define patterns to match (without price parts)
    patterns = {
        "השלם לסט:הוסף מכנס+גרביים": "+ Shorts\n+ Socks",
        "השלם לסט:הוסף מכנס": "+ Shorts",
        "הוסף מכנסיים:כן": "+ Pants",
        "הוסף גרביים:כן": "+ Socks"
    }
    
    # Check each pattern
    for pattern, result in patterns.items():
        # Look for the pattern at the start of the field (before any price)
        if pattern in options_text:
            return result
    
    # If no pattern matches, return empty
    return ""


def extract_name(notes_text):
    """Extract name from notes (שם מאחור field)."""
    if not notes_text:
        return ""
    
    # Try Hebrew pattern
    name = extract_option_value(notes_text, "שם מאחור באנגלית")
    if name:
        return name
    
    # Try other patterns
    name = extract_option_value(notes_text, "שם מאחור")
    return name


def extract_number(notes_text):
    """Extract number from notes (מספר מאחור field)."""
    if not notes_text:
        return ""
    
    return extract_option_value(notes_text, "מספר מאחור")


def parse_shipping_label_simple(shipping_label):
    """
    Parse shipping label and format it without translation.
    Format: "Name / Street / City / State / Country / Phone"
    
    Returns formatted address string.
    """
    if not shipping_label or not shipping_label.strip():
        return None
    
    parts = [p.strip() for p in shipping_label.split('/')]
    
    # Expected format: Name / Street / City / State / Country / Phone
    name = parts[0] if len(parts) > 0 else ''
    street = parts[1] if len(parts) > 1 else ''
    city = parts[2] if len(parts) > 2 else ''
    zip_code = parts[3] if len(parts) > 3 else ''
    country = parts[4] if len(parts) > 4 else ''
    phone = parts[5] if len(parts) > 5 else ''
    
    # Remove quotes from phone
    phone = phone.strip('"').strip()
    
    # Format the address
    formatted_lines = []
    formatted_lines.append(f"Name: {name if name else 'N/A'}")
    formatted_lines.append(f"Address: {street if street else 'N/A'}")
    formatted_lines.append(f"City: {city if city else 'N/A'}")
    formatted_lines.append(f"Zip Code: {zip_code if zip_code else 'N/A'}")
    formatted_lines.append(f"Phone: {phone if phone else 'N/A'}")
    
    return '\n'.join(formatted_lines)


def is_settlement_type_city(city_name):
    """
    Check if city name indicates it's a settlement, kibbutz, moshav, or village.
    These types typically don't have street addresses.
    """
    if not city_name:
        return False
    
    city_lower = city_name.lower()
    
    # Hebrew keywords
    settlement_keywords = [
        'התיישבות', 'יישוב',  # settlement
        'קיבוץ', 'קיבוץ',      # kibbutz
        'מושב',                 # moshav
        'כפר',                  # village/kfar
        'מצפה',                 # lookout point
        'גבעת',                 # hill
        'נווה',                 # dwelling
    ]
    
    # English keywords
    settlement_keywords_en = [
        'kibbutz', 'moshav', 'kfar', 'settlement', 'village'
    ]
    
    all_keywords = settlement_keywords + settlement_keywords_en
    
    for keyword in all_keywords:
        if keyword in city_lower:
            return True
    
    return False


def format_address_with_claude(shipping_label, order_no):
    """
    Use Claude API to format and translate address.
    Falls back to Hebrew formatting if API is not available.
    
    Args:
        shipping_label: Raw shipping label string (e.g., "Name / Street / City / / Country / Phone")
        order_no: Order number for tracking
    
    Returns:
        Dictionary with formatted address info and status
    """
    if not shipping_label or not shipping_label.strip():
        return {
            'formatted': '',
            'missing_fields': ['all'],
            'completed_by_script': False,
            'raw': shipping_label
        }
    
    api_key = get_api_key()
    
    if not api_key:
        # Fallback: Format without translation
        print(f"Info: Formatting address for order {order_no} without translation (no API key)")
        formatted = parse_shipping_label_simple(shipping_label)
        
        if formatted:
            missing_fields = []
            for line in formatted.split('\n'):
                if 'N/A' in line:
                    field_name = line.split(':')[0].strip()
                    missing_fields.append(field_name)
            
            return {
                'formatted': formatted,
                'missing_fields': missing_fields,
                'completed_by_script': False,
                'raw': shipping_label
            }
        else:
            return {
                'formatted': shipping_label,
                'missing_fields': ['parse_error'],
                'completed_by_script': False,
                'raw': shipping_label
            }
    
    prompt = f"""Given this shipping label from an order, please format it into a structured address with the following format:

Name: [Full name in English]
Address: [Street address in English]
City: [City name in English]
Zip Code: [Postal code]
Phone: [Phone number]

Shipping label: {shipping_label}

Requirements:
1. Translate any Hebrew text to English
2. If the zip code is missing (empty field between City and Country), look it up based on the street and city in Israel
3. Keep the exact format with "Name:", "Address:", "City:", "Zip Code:", "Phone:" labels
4. Each field on a new line
5. If a field is completely missing from the input, write "N/A" for that field
6. Remove any quotes from phone numbers

Please provide ONLY the formatted address, nothing else."""

    try:
        response = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "Content-Type": "application/json",
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01"
            },
            json={
                "model": "claude-sonnet-4-20250514",
                "max_tokens": 1000,
                "messages": [
                    {"role": "user", "content": prompt}
                ]
            },
            timeout=30
        )
        
        if response.status_code == 200:
            data = response.json()
            formatted_address = data['content'][0]['text'].strip()
            
            # Analyze what fields are present
            missing_fields = []
            completed_zip = False
            
            lines = formatted_address.split('\n')
            for line in lines:
                if 'N/A' in line:
                    field_name = line.split(':')[0].strip()
                    missing_fields.append(field_name)
            
            # Check if zip code was added (not in original but present in formatted)
            if 'Zip Code: N/A' not in formatted_address and ('/ /' in shipping_label or '///' in shipping_label):
                # Zip might have been added by Claude
                completed_zip = True
            
            return {
                'formatted': formatted_address,
                'missing_fields': missing_fields,
                'completed_by_script': completed_zip,
                'raw': shipping_label
            }
        elif response.status_code == 401:
            # Invalid API key
            print(f"Error: Invalid API key (401 Unauthorized)")
            clear_invalid_api_key()
            print("Please run the script again to enter a new API key.\n")
            
            # Fallback to Hebrew formatting
            formatted = parse_shipping_label_simple(shipping_label)
            if formatted:
                missing_fields = []
                for line in formatted.split('\n'):
                    if 'N/A' in line:
                        field_name = line.split(':')[0].strip()
                        missing_fields.append(field_name)
                
                return {
                    'formatted': formatted,
                    'missing_fields': missing_fields,
                    'completed_by_script': False,
                    'raw': shipping_label
                }
            else:
                return {
                    'formatted': shipping_label,
                    'missing_fields': ['api_key_invalid'],
                    'completed_by_script': False,
                    'raw': shipping_label
                }
        else:
            print(f"Warning: Claude API error for order {order_no}: {response.status_code}")
            # Fallback to Hebrew formatting
            formatted = parse_shipping_label_simple(shipping_label)
            if formatted:
                missing_fields = []
                for line in formatted.split('\n'):
                    if 'N/A' in line:
                        field_name = line.split(':')[0].strip()
                        missing_fields.append(field_name)
                
                return {
                    'formatted': formatted,
                    'missing_fields': missing_fields,
                    'completed_by_script': False,
                    'raw': shipping_label
                }
            else:
                return {
                    'formatted': shipping_label,
                    'missing_fields': ['api_error'],
                    'completed_by_script': False,
                    'raw': shipping_label
                }
    
    except Exception as e:
        print(f"Warning: Error formatting address for order {order_no}: {str(e)}")
        # Fallback to Hebrew formatting
        formatted = parse_shipping_label_simple(shipping_label)
        if formatted:
            missing_fields = []
            for line in formatted.split('\n'):
                if 'N/A' in line:
                    field_name = line.split(':')[0].strip()
                    missing_fields.append(field_name)
            
            return {
                'formatted': formatted,
                'missing_fields': missing_fields,
                'completed_by_script': False,
                'raw': shipping_label
            }
        else:
            return {
                'formatted': shipping_label,
                'missing_fields': ['processing_error'],
                'completed_by_script': False,
                'raw': shipping_label
            }


def extract_address(row, order_no, translations):
    """Extract and format address from shipping information using Claude API."""
    shipping_label = row.get('Shipping label', '').strip()
    
    if shipping_label:
        # Use Claude API to format the address (or fallback to Hebrew)
        address_info = format_address_with_claude(shipping_label, order_no)
        
        # Check if address is missing and city is a settlement type
        if address_info.get('missing_fields') and 'Address' in address_info.get('missing_fields', []):
            # Try to extract city from the formatted address or raw label
            city = ''
            if address_info['formatted']:
                for line in address_info['formatted'].split('\n'):
                    if line.startswith('City:'):
                        city = line.split(':', 1)[1].strip()
                        if city != 'N/A':
                            break
            
            # If no city found in formatted, try parsing raw
            if not city or city == 'N/A':
                parts = shipping_label.split('/')
                if len(parts) > 2:
                    city = parts[2].strip()
            
            # Check if it's a settlement type
            if city and is_settlement_type_city(city):
                # Replace "Address: N/A" with "Address: [City Name]"
                formatted_lines = address_info['formatted'].split('\n')
                for i, line in enumerate(formatted_lines):
                    if line.startswith('Address:') and 'N/A' in line:
                        formatted_lines[i] = f"Address: {city}"
                        # Remove 'Address' from missing fields
                        if 'Address' in address_info['missing_fields']:
                            address_info['missing_fields'].remove('Address')
                        break
                
                address_info['formatted'] = '\n'.join(formatted_lines)
        
        return address_info
    
    # Fallback to old method if no shipping label
    street = row.get('Delivery address', '').strip()
    city = row.get('Delivery city', '').strip()
    
    if street and city:
        simple_address = f"{street}, {city}"
    elif city:
        # If only city and it's a settlement type, use city as address
        if is_settlement_type_city(city):
            simple_address = f"Address: {city}\nCity: {city}"
        else:
            simple_address = city
    elif street:
        simple_address = street
    else:
        simple_address = ""
    
    return {
        'formatted': simple_address,
        'missing_fields': ['formatted_address_not_available'],
        'completed_by_script': False,
        'raw': simple_address
    }


def process_order(row, translations):
    """Process a single order row and extract relevant information."""
    order_no = row.get('Order number', '').strip()
    product = row.get('Item', '').strip()
    quantity = int(row.get('Qty', '1'))
    options = row.get('Variant', '')
    notes = row.get('Custom text', '')
    
    # Extract details
    size = extract_size(options)
    patch = extract_patch(options, product, translations)  # Pass product name for team extraction
    adds = extract_adds(options)
    name = extract_name(notes)
    number = extract_number(notes)
    address_info = extract_address(row, order_no, translations)  # Pass translations for caching
    
    # Find product image
    product_image = find_product_image(product, translations)
    
    # Create order details
    order_details = {
        'order_no': order_no,
        'product': product,
        'product_image': product_image,  # Store image path
        'size': size,
        'name': name,
        'number': number,
        'patch': patch,
        'adds': adds,
        'address': address_info['formatted'],  # Use formatted address
        'address_info': address_info,  # Keep full info for reporting
        'quantity': quantity
    }
    
    return order_details


def create_word_table(orders, output_file='orders_table.docx'):
    """Create a Word document with a table of orders."""
    doc = Document()
    
    # Set up document properties
    section = doc.sections[0]
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # Add title
    title = doc.add_heading('Orders Summary', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Create table with header
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Light Grid Accent 1'
    
    # Set column headers
    headers = ['Order no.', 'Product', 'Size', 'Name', 'Number', 'Patch', 'Adds', 'Address']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Bold header text
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    # Group orders by order number to handle multi-product orders
    # Format: {order_no: [order1, order2, ...]}
    grouped_orders = {}
    for order in orders:
        order_no = order['order_no']
        if order_no not in grouped_orders:
            grouped_orders[order_no] = []
        # Expand by quantity - each quantity becomes a separate row
        for _ in range(order['quantity']):
            grouped_orders[order_no].append(order)
    
    # Track merge information for each order number
    # Format: {order_no: {'start_row': int, 'row_count': int}}
    merge_info = {}
    current_row = 1  # Start after header
    
    # Add all rows and track merge requirements
    for order_no in sorted(grouped_orders.keys()):
        order_items = grouped_orders[order_no]
        row_count = len(order_items)
        
        # Track if this order needs merging (more than one row)
        if row_count > 1:
            merge_info[order_no] = {
                'start_row': current_row,
                'row_count': row_count
            }
        
        # Add a row for each item (already expanded by quantity)
        for order in order_items:
            row_cells = table.add_row().cells
            row_cells[0].text = order['order_no']
            
            # Product column - insert image if available, otherwise text
            if order.get('product_image') and os.path.exists(order['product_image']):
                success = insert_image_in_cell(row_cells[1], order['product_image'], width_inches=2.0)
                if not success:
                    # Fallback to text if image insertion fails
                    row_cells[1].text = order['product']
            else:
                # No image found - use product name
                row_cells[1].text = order['product']
            
            row_cells[2].text = order['size']
            row_cells[3].text = order['name']
            row_cells[4].text = order['number']
            row_cells[5].text = order['patch']
            row_cells[6].text = order['adds']
            row_cells[7].text = order['address']
            
            # Set font size for data (skip Product column as it may have image)
            for idx, cell in enumerate(row_cells):
                if idx == 1:  # Skip product column
                    continue
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
            
            current_row += 1
    
    # Apply vertical merges for multi-product orders
    for order_no, info in merge_info.items():
        start_row = info['start_row']
        row_count = info['row_count']
        
        # Merge Order No. column (column 0)
        for i in range(row_count):
            row_idx = start_row + i
            cell = table.rows[row_idx].cells[0]
            if i == 0:
                set_cell_vertical_merge(cell, 'restart')
            else:
                set_cell_vertical_merge(cell, 'continue')
        
        # Merge Address column (column 7)
        for i in range(row_count):
            row_idx = start_row + i
            cell = table.rows[row_idx].cells[7]
            if i == 0:
                set_cell_vertical_merge(cell, 'restart')
            else:
                set_cell_vertical_merge(cell, 'continue')
    
    # Adjust column widths
    widths = [Inches(0.8), Inches(2.5), Inches(0.6), Inches(1.2), 
              Inches(0.7), Inches(1.2), Inches(0.8), Inches(1.7)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    
    # Save document
    doc.save(output_file)
    return output_file


def main():
    # Parse command line arguments
    import argparse
    parser = argparse.ArgumentParser(description='Process orders CSV and generate Word document')
    parser.add_argument('csv_file', nargs='?', help='Path to CSV file (optional, will use latest if not provided)')
    parser.add_argument('--no-translate', action='store_true', help='Skip address translation (use Hebrew addresses)')
    args = parser.parse_args()
    
    # Get CSV file path
    if args.csv_file:
        csv_file = args.csv_file
    else:
        csv_file = find_latest_csv()
        if not csv_file:
            print("Error: No CSV file found in current directory")
            print("Usage: python process_orders.py [csv_file] [--no-translate]")
            sys.exit(1)
        print(f"Using CSV file: {csv_file}")
    
    if not os.path.exists(csv_file):
        print(f"Error: File '{csv_file}' not found")
        sys.exit(1)
    
    # Set global flag for translation
    if args.no_translate:
        global NO_TRANSLATE
        NO_TRANSLATE = True
        print("Translation disabled - using Hebrew addresses\n")
    
    # Load translations
    print("Loading translations...")
    translations = load_translations()
    print(f"Loaded {len(translations)} translation sections")
    
    # Read and process CSV
    orders = []
    with open(csv_file, 'r', encoding='utf-8-sig') as f:  # utf-8-sig handles BOM
        reader = csv.DictReader(f)
        for row in reader:
            order = process_order(row, translations)
            orders.append(order)
    
    if not orders:
        print("No orders found in CSV file")
        sys.exit(1)
    
    print(f"\nProcessed {len(orders)} orders")
    total_items = sum(order['quantity'] for order in orders)
    print(f"Total items: {total_items}")
    
    # Create Word document
    current_date = datetime.now().strftime('%Y-%m-%d')
    output_file = f'orders_{current_date}.docx'
    created_file = create_word_table(orders, output_file)
    print(f"\nWord document created: {created_file}")
    
    # Address validation report
    print("\n" + "="*80)
    print("ADDRESS VALIDATION REPORT")
    print("="*80)
    
    orders_with_issues = []
    orders_with_completions = []
    
    for order in orders:
        address_info = order['address_info']
        missing_fields = address_info['missing_fields']
        completed = address_info['completed_by_script']
        
        if missing_fields and missing_fields != []:
            if missing_fields not in [['api_error'], ['processing_error'], ['formatted_address_not_available']]:
                orders_with_issues.append({
                    'order_no': order['order_no'],
                    'missing_fields': missing_fields,
                    'raw_address': address_info['raw']
                })
        
        if completed:
            orders_with_completions.append({
                'order_no': order['order_no'],
                'formatted_address': order['address']
            })
    
    if not orders_with_issues and not orders_with_completions:
        print("\n✓ All addresses are complete with no missing fields!")
    
    if orders_with_issues:
        print(f"\n⚠ Orders with missing address fields ({len(orders_with_issues)}):")
        print("-" * 80)
        for issue in orders_with_issues:
            print(f"\nOrder {issue['order_no']}:")
            print(f"  Missing fields: {', '.join(issue['missing_fields'])}")
            print(f"  Raw address: {issue['raw_address'][:100]}...")
    
    if orders_with_completions:
        print(f"\n✓ Addresses completed by script ({len(orders_with_completions)}):")
        print("-" * 80)
        for completion in orders_with_completions:
            print(f"\nOrder {completion['order_no']}:")
            print(f"  Completed address:")
            for line in completion['formatted_address'].split('\n'):
                print(f"    {line}")
    
    print("\n" + "="*80)


if __name__ == '__main__':
    main()