#!/usr/bin/env python3
"""
Process orders CSV file and generate a Word document with order details.
Each order is represented as a row with: Order no., Product, Size, Name, Number, Patch, Adds, Address
Uses translations.json for field translations (e.g., Hebrew to English).
"""

import csv
import sys
import os
import re
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import glob


def set_cell_vertical_merge(cell, merge_type):
    """
    Set vertical merge for a cell.
    
    Args:
        cell: The cell to set merge on
        merge_type: 'restart' for first cell in merge, 'continue' for subsequent cells
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    vMerge = OxmlElement('w:vMerge')
    if merge_type == 'restart':
        vMerge.set(qn('w:val'), 'restart')
    # For 'continue', don't set the val attribute (empty vMerge element)
    tcPr.append(vMerge)


# Translation file path
TRANSLATION_FILE = 'translations.json'


def load_translations():
    """Load translations from JSON file."""
    if os.path.exists(TRANSLATION_FILE):
        with open(TRANSLATION_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    # Return default structure with team_leagues section
    return {
        "patch": {},
        "team_leagues": {}
    }


def save_translations(translations):
    """Save translations to JSON file."""
    with open(TRANSLATION_FILE, 'w', encoding='utf-8') as f:
        json.dump(translations, f, ensure_ascii=False, indent=2)


def translate_field(field_name, value, translations):
    """
    Translate a field value using the translations dictionary.
    If the field section exists but the value is not found, prompt user for translation.
    
    Args:
        field_name: The field section name (e.g., 'patch')
        value: The value to translate
        translations: The translations dictionary
    
    Returns:
        Translated value or original if no translation needed
    """
    if not value or not value.strip():
        return value
    
    value = value.strip()
    
    # Check if this field has a translation section
    if field_name not in translations:
        return value
    
    field_translations = translations[field_name]
    
    # Check if translation exists
    if value in field_translations:
        return field_translations[value]
    
    # Translation doesn't exist - prompt user
    print(f"\n{'='*60}")
    print(f"Translation needed for field: '{field_name}'")
    print(f"Hebrew text: {value}")
    user_translation = input(f"Enter English translation (or press Enter to keep as-is): ").strip()
    print(f"{'='*60}\n")
    
    # Save the translation
    if user_translation:
        translations[field_name][value] = user_translation
        save_translations(translations)
        return user_translation
    else:
        # User chose to keep as-is, save this choice
        translations[field_name][value] = value
        save_translations(translations)
        return value


def extract_team_name(product_name):
    """
    Extract team name from product name.
    Pattern: text between [חולצת|חולצה ארוכה|חולצת נשים|חליפת ילדים|מכנס] and [בית|חוץ|השלישי|השלישית]
    
    Args:
        product_name: The product name string
    
    Returns:
        Team name if found, None otherwise
    """
    if not product_name:
        return None
    
    # Define the prefix and suffix patterns
    prefix_pattern = r'(?:חולצת|חולצה ארוכה|חולצת נשים|חליפת ילדים|מכנס)\s+'
    suffix_pattern = r'\s+(?:בית|חוץ|השלישי|השלישית)'
    
    # Create full pattern to extract team name
    pattern = prefix_pattern + r'(.+?)' + suffix_pattern
    
    match = re.search(pattern, product_name)
    if match:
        return match.group(1).strip()
    
    return None


def get_league_patch_for_team(team_name, translations):
    """
    Get the league patch for a specific team.
    
    Args:
        team_name: The team name extracted from product
        translations: The translations dictionary
    
    Returns:
        League patch string (e.g., "LFP", "Premier League", etc.)
    """
    # Ensure team_leagues section exists
    if "team_leagues" not in translations:
        translations["team_leagues"] = {}
    
    team_leagues = translations["team_leagues"]
    
    # Check if team exists in mapping
    if team_name in team_leagues:
        return team_leagues[team_name]
    
    # Team not found - prompt user
    print(f"\n{'='*60}")
    print(f"League patch needed for team: '{team_name}'")
    print(f"Product contains team '{team_name}' but no league mapping exists.")
    user_patch = input(f"Enter league patch for this team (or press Enter for default 'League'): ").strip()
    print(f"{'='*60}\n")
    
    # Save the mapping
    if user_patch:
        translations["team_leagues"][team_name] = user_patch
        save_translations(translations)
        return user_patch
    else:
        # User chose default, save this choice
        translations["team_leagues"][team_name] = "League"
        save_translations(translations)
        return "League"


def find_latest_csv(directory='.'):
    """Find the most recent CSV file in the directory."""
    csv_files = glob.glob(os.path.join(directory, '*.csv'))
    if not csv_files:
        return None
    return max(csv_files, key=os.path.getctime)


def extract_option_value(options_text, key):
    """Extract value for a specific key from options string."""
    if not options_text:
        return ""
    
    # Pattern to match "key:value" where value can be in Hebrew or English
    pattern = rf'{re.escape(key)}:([^|]+)'
    match = re.search(pattern, options_text)
    if match:
        return match.group(1).strip()
    return ""


def extract_size(options_text):
    """Extract size from options string (מידה field)."""
    size = extract_option_value(options_text, "מידה")
    if size:
        # Extract just the size number/text before the colon or description
        size_match = re.match(r'([^:]+)', size)
        if size_match:
            return size_match.group(1).strip()
    return ""


def extract_patch(options_text, product_name, translations):
    """
    Extract patch information from options string (פאץ׳ field) and translate it.
    Special handling for "ליגה" - extract team name and lookup corresponding league patch.
    
    Args:
        options_text: The options/variant text containing patch info
        product_name: The product name for team extraction
        translations: The translations dictionary
    
    Returns:
        Translated patch string
    """
    patch = extract_option_value(options_text, "פאץ'") or extract_option_value(options_text, "פאץ׳")
    
    if not patch:
        return ""
    
    # Special handling for "ליגה"
    if patch == "ליגה":
        # Try to extract team name from product
        team_name = extract_team_name(product_name)
        
        if team_name:
            # Get league patch for this team
            return get_league_patch_for_team(team_name, translations)
        else:
            # Pattern doesn't match, return default "League"
            return "League"
    
    # For other patches, use normal translation
    return translate_field('patch', patch, translations)


def extract_adds(options_text):
    """
    Extract additional items from options string (הוסף field).
    Handles various add-on patterns and returns formatted English text.
    
    Patterns:
    - השלם לסט:הוסף מכנס -> + Shorts
    - השלם לסט:הוסף מכנס+גרביים -> + Shorts\n+ Socks
    - הוסף גרביים:כן -> + Socks
    - הוסף מכנסיים:כן -> + Pants
    """
    if not options_text:
        return ""
    
    # Define patterns to match (without price parts)
    patterns = {
        "השלם לסט:הוסף מכנס+גרביים": "+ Shorts\n+ Socks",
        "השלם לסט:הוסף מכנס": "+ Shorts",
        "הוסף מכנסיים:כן": "+ Pants",
        "הוסף גרביים:כן": "+ Socks"
    }
    
    # Check each pattern
    for pattern, result in patterns.items():
        # Look for the pattern at the start of the field (before any price)
        if pattern in options_text:
            return result
    
    # If no pattern matches, return empty
    return ""


def extract_name(notes_text):
    """Extract name from notes (שם מאחור field)."""
    if not notes_text:
        return ""
    
    # Try Hebrew pattern
    name = extract_option_value(notes_text, "שם מאחור באנגלית")
    if name:
        return name
    
    # Try other patterns
    name = extract_option_value(notes_text, "שם מאחור")
    return name


def extract_number(notes_text):
    """Extract number from notes (מספר מאחור field)."""
    if not notes_text:
        return ""
    
    return extract_option_value(notes_text, "מספר מאחור")


def extract_address(row):
    """Extract address from shipping information."""
    # Try to get address from the combined address field first
    if row.get('Shipping label') and row['Shipping label'].strip():
        # Parse the combined address format: "Name / Street / City / State / Country / Phone"
        parts = [p.strip() for p in row['Shipping label'].split('/')]
        if len(parts) >= 3:
            street = parts[1] if len(parts) > 1 else ""
            city = parts[2] if len(parts) > 2 else ""
            return f"{street}, {city}".strip(', ')
    
    # Fallback to individual fields
    street = row.get('Delivery address', '').strip()
    city = row.get('Delivery city', '').strip()
    
    if street and city:
        return f"{street}, {city}"
    elif street:
        return street
    elif city:
        return city
    
    return ""


def process_order(row, translations):
    """Process a single order row and extract relevant information."""
    order_no = row.get('Order number', '').strip()
    product = row.get('Item', '').strip()
    quantity = int(row.get('Qty', '1'))
    options = row.get('Variant', '')
    notes = row.get('Custom text', '')
    
    # Extract details
    size = extract_size(options)
    patch = extract_patch(options, product, translations)  # Pass product name for team extraction
    adds = extract_adds(options)
    name = extract_name(notes)
    number = extract_number(notes)
    address = extract_address(row)
    
    # Create order details
    order_details = {
        'order_no': order_no,
        'product': product,
        'size': size,
        'name': name,
        'number': number,
        'patch': patch,
        'adds': adds,
        'address': address,
        'quantity': quantity
    }
    
    return order_details


def create_word_table(orders, output_file='orders_table.docx'):
    """Create a Word document with a table of orders."""
    doc = Document()
    
    # Set up document properties
    section = doc.sections[0]
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # Add title
    title = doc.add_heading('Orders Summary', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Create table with header
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Light Grid Accent 1'
    
    # Set column headers
    headers = ['Order no.', 'Product', 'Size', 'Name', 'Number', 'Patch', 'Adds', 'Address']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Bold header text
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    # Group orders by order number to handle multi-product orders
    # Format: {order_no: [order1, order2, ...]}
    grouped_orders = {}
    for order in orders:
        order_no = order['order_no']
        if order_no not in grouped_orders:
            grouped_orders[order_no] = []
        # Expand by quantity - each quantity becomes a separate row
        for _ in range(order['quantity']):
            grouped_orders[order_no].append(order)
    
    # Track merge information for each order number
    # Format: {order_no: {'start_row': int, 'row_count': int}}
    merge_info = {}
    current_row = 1  # Start after header
    
    # Add all rows and track merge requirements
    for order_no in sorted(grouped_orders.keys()):
        order_items = grouped_orders[order_no]
        row_count = len(order_items)
        
        # Track if this order needs merging (more than one row)
        if row_count > 1:
            merge_info[order_no] = {
                'start_row': current_row,
                'row_count': row_count
            }
        
        # Add a row for each item (already expanded by quantity)
        for order in order_items:
            row_cells = table.add_row().cells
            row_cells[0].text = order['order_no']
            row_cells[1].text = order['product']
            row_cells[2].text = order['size']
            row_cells[3].text = order['name']
            row_cells[4].text = order['number']
            row_cells[5].text = order['patch']
            row_cells[6].text = order['adds']
            row_cells[7].text = order['address']
            
            # Set font size for data
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
            
            current_row += 1
    
    # Apply vertical merges for multi-product orders
    for order_no, info in merge_info.items():
        start_row = info['start_row']
        row_count = info['row_count']
        
        # Merge Order No. column (column 0)
        for i in range(row_count):
            row_idx = start_row + i
            cell = table.rows[row_idx].cells[0]
            if i == 0:
                set_cell_vertical_merge(cell, 'restart')
            else:
                set_cell_vertical_merge(cell, 'continue')
        
        # Merge Address column (column 7)
        for i in range(row_count):
            row_idx = start_row + i
            cell = table.rows[row_idx].cells[7]
            if i == 0:
                set_cell_vertical_merge(cell, 'restart')
            else:
                set_cell_vertical_merge(cell, 'continue')
    
    # Adjust column widths
    widths = [Inches(0.8), Inches(2.5), Inches(0.6), Inches(1.2), 
              Inches(0.7), Inches(1.2), Inches(0.8), Inches(1.7)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    
    # Save document
    doc.save(output_file)
    return output_file


def main():
    # Get CSV file path
    if len(sys.argv) > 1:
        csv_file = sys.argv[1]
    else:
        csv_file = find_latest_csv()
        if not csv_file:
            print("Error: No CSV file found in current directory")
            print("Usage: python process_orders.py [csv_file]")
            sys.exit(1)
        print(f"Using CSV file: {csv_file}")
    
    if not os.path.exists(csv_file):
        print(f"Error: File '{csv_file}' not found")
        sys.exit(1)
    
    # Load translations
    print("Loading translations...")
    translations = load_translations()
    print(f"Loaded {len(translations)} translation sections")
    
    # Read and process CSV
    orders = []
    with open(csv_file, 'r', encoding='utf-8-sig') as f:  # utf-8-sig handles BOM
        reader = csv.DictReader(f)
        for row in reader:
            order = process_order(row, translations)
            orders.append(order)
    
    if not orders:
        print("No orders found in CSV file")
        sys.exit(1)
    
    print(f"\nProcessed {len(orders)} orders")
    total_items = sum(order['quantity'] for order in orders)
    print(f"Total items: {total_items}")
    
    # Create Word document
    output_file = 'orders_table.docx'
    created_file = create_word_table(orders, output_file)
    print(f"\nWord document created: {created_file}")
    
    # Print summary of first order as example
    if orders:
        print("\nExample (first order):")
        first = orders[0]
        print(f"  Order no.: {first['order_no']}")
        print(f"  Product: {first['product']}")
        print(f"  Size: {first['size']}")
        print(f"  Name: {first['name']}")
        print(f"  Number: {first['number']}")
        print(f"  Patch: {first['patch']}")
        print(f"  Adds: {first['adds']}")
        print(f"  Address: {first['address']}")


if __name__ == '__main__':
    main()